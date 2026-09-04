import io
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_PATH = Path("modelo/Evidencia de Testes - CORSAN - Modelo.docx")

QA_NAMES = ["Victor Morbach", "Aline Rodrigues Vieira Pinto"]

STATUS_OPCOES = ["Concluído", "Em andamento", "Bloqueado"]

SHEET_NAME = "plano de teste"
COL_PROCESSO = "Processos"
COL_CASO_TESTE = "Caso de Teste"

IMAGEM_LARGURA_POL = 6.0

st.set_page_config(page_title="Gerador de Evidência de Testes", layout="centered")
st.title("Gerador de Evidência de Testes - CORSAN")


# ------------------------------------------------------------------
# FUNÇÕES AUXILIARES
# ------------------------------------------------------------------
def get_unique_cells(row):
    """Remove células repetidas causadas por merge horizontal, mantendo a ordem."""
    unique = []
    seen_ids = set()
    for cell in row.cells:
        if id(cell._tc) not in seen_ids:
            seen_ids.add(id(cell._tc))
            unique.append(cell)
    return unique


def set_cell_text(cell, texto):
    """Escreve o texto na célula preservando a formatação do primeiro run, se existir."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.text = texto
        for run_extra in cell.paragraphs[0].runs[1:]:
            run_extra.text = ""
    else:
        cell.text = texto


def sanitize_filename(nome):
    """Remove caracteres inválidos para nomes de arquivo no Windows."""
    return re.sub(r'[\\/:*?"<>|]', "", nome).strip()


def fill_table(doc, valores):
    """Percorre a tabela do modelo e preenche o valor ao lado de cada rótulo conhecido."""
    table = doc.tables[0]
    for row in table.rows:
        cells = get_unique_cells(row)
        for i, cell in enumerate(cells):
            rotulo = cell.text.strip().upper()
            if rotulo in valores and i + 1 < len(cells):
                set_cell_text(cells[i + 1], valores[rotulo])
    return table


def set_status(doc, status_valor):
    """Localiza o parágrafo 'STATUS:' e adiciona o valor escolhido ao lado."""
    for paragraph in doc.paragraphs:
        if "STATUS:" in paragraph.text.upper():
            run = paragraph.add_run(f" {status_valor}")
            run.bold = True
            return paragraph
    return None


def add_screenshots_after(doc, anchor_element, screenshots):
    """Insere os prints logo após o elemento âncora (ex.: parágrafo do STATUS),
    sem quebra de página, mantendo a ordem de upload.
    A legenda fica logo abaixo da imagem correspondente."""
    current_anchor = anchor_element
    for idx, item in enumerate(screenshots, start=1):
        arquivo = item["arquivo"]
        legenda = item.get("legenda", "")

        # 1) Imagem primeiro
        p_imagem = doc.add_paragraph()
        p_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_imagem = p_imagem.add_run()
        arquivo.seek(0)
        run_imagem.add_picture(arquivo, width=Inches(IMAGEM_LARGURA_POL))

        current_anchor.addnext(p_imagem._p)
        current_anchor = p_imagem._p

        # 2) Legenda logo abaixo da imagem
        p_legenda = doc.add_paragraph()
        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_legenda = p_legenda.add_run(
            f"Print {idx} - {legenda}" if legenda else f"Print {idx}"
        )
        run_legenda.italic = True
        run_legenda.font.size = Pt(10)

        current_anchor.addnext(p_legenda._p)
        current_anchor = p_legenda._p

    return doc


def fill_document(template_bytes, valores, status_valor, screenshots):
    doc = Document(io.BytesIO(template_bytes))

    valores_upper = {k.upper(): v for k, v in valores.items()}
    fill_table(doc, valores_upper)

    status_paragraph = set_status(doc, status_valor)
    anchor = status_paragraph._p if status_paragraph is not None else doc.tables[0]._tbl

    if screenshots:
        add_screenshots_after(doc, anchor, screenshots)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ------------------------------------------------------------------
# INTERFACE
# ------------------------------------------------------------------
if not TEMPLATE_PATH.exists():
    st.error(f"Modelo não encontrado em: {TEMPLATE_PATH}")
    st.stop()

with open(TEMPLATE_PATH, "rb") as f:
    template_bytes = f.read()

arquivo_xlsx = st.file_uploader("Envie o Caderno de Testes (.xlsx)", type=["xlsx"])

if arquivo_xlsx:
    xls = pd.ExcelFile(arquivo_xlsx)

    if SHEET_NAME in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=SHEET_NAME)
    else:
        st.warning(
            f"Aba '{SHEET_NAME}' não encontrada. Usando a primeira aba: '{xls.sheet_names[0]}'."
        )
        df = pd.read_excel(xls, sheet_name=xls.sheet_names[0])

    # ---------------- Cenário ----------------
    cenarios = df[COL_PROCESSO].dropna().unique().tolist()
    if len(cenarios) == 1:
        cenario = cenarios[0]
        st.write(f"**Cenário:** {cenario}")
    else:
        cenario = st.selectbox("Cenário", cenarios)

    # ---------------- Caso de Teste (filtrado pelo Cenário) ----------------
    df_filtrado = df[df[COL_PROCESSO] == cenario]
    casos_teste = df_filtrado[COL_CASO_TESTE].dropna().unique().tolist()
    if len(casos_teste) == 1:
        caso_teste = casos_teste[0]
        st.write(f"**Caso de Teste:** {caso_teste}")
    else:
        caso_teste = st.selectbox(
            "Caso de Teste", casos_teste, key=f"caso_teste_select_{cenario}"
        )

    # ---------------- QA ----------------
    if len(QA_NAMES) == 1:
        qa = QA_NAMES[0]
        st.write(f"**QA:** {qa}")
    else:
        qa = st.selectbox("QA", QA_NAMES)

    # ---------------- DEV / IPC / Squad ----------------
    dev = st.text_input("DEV")
    ipc = st.text_input("IPC")
    squad = st.text_input("Squad", value="CORSAN")

    # ---------------- Status ----------------
    status = st.selectbox("Status", STATUS_OPCOES)

    st.divider()
    st.subheader("Prints de tela")

    if "num_prints" not in st.session_state:
        st.session_state.num_prints = 1

    if "uploader_version" not in st.session_state:
        st.session_state.uploader_version = 0

    v = st.session_state.uploader_version

    col_add, col_remove = st.columns(2)
    with col_add:
        if st.button("+ Adicionar print"):
            st.session_state.num_prints += 1
    with col_remove:
        if st.session_state.num_prints > 1 and st.button("- Remover último print"):
            st.session_state.num_prints -= 1

    screenshots = []
    for i in range(st.session_state.num_prints):
        st.markdown(f"**Print {i + 1}**")
        arquivo_print = st.file_uploader(
            f"Imagem {i + 1}",
            type=["png", "jpg", "jpeg"],
            key=f"print_{i}_{v}",
        )
        legenda_print = st.text_input(
            "Legenda (opcional)",
            key=f"legenda_{i}_{v}",
        )
        if arquivo_print is not None:
            st.image(arquivo_print, width=250)
            screenshots.append({"arquivo": arquivo_print, "legenda": legenda_print})

    st.divider()

    if st.button("Gerar documento"):
        valores = {
            "Cenário": cenario,
            "Caso de Teste": caso_teste,
            "QA": qa,
            "DEV": dev,
            "IPC": ipc,
            "Squad": squad,
        }

        resultado = fill_document(template_bytes, valores, status, screenshots)

        nome_arquivo = sanitize_filename(
            f"Evidencias de teste {squad} {cenario} {caso_teste}"
        ) + ".docx"

        st.session_state["documento_gerado"] = resultado.getvalue()
        st.session_state["nome_arquivo_gerado"] = nome_arquivo

    if "documento_gerado" in st.session_state:
        st.success("Documento gerado com sucesso.")
        baixou = st.download_button(
            label="Baixar documento",
            data=st.session_state["documento_gerado"],
            file_name=st.session_state["nome_arquivo_gerado"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="botao_download",
        )

        if baixou:
            st.session_state.num_prints = 1
            st.session_state.uploader_version += 1
            del st.session_state["documento_gerado"]
            del st.session_state["nome_arquivo_gerado"]
            st.rerun()
else:
    st.info("Envie o arquivo .xlsx do Caderno de Testes para começar.")
